"""
cleaner.py
----------
Reusable data-cleaning logic for NeatSheet.
This is the same cleaning logic from the original script,
just packaged as a function so the web app (app.py) can call it.
"""

import pandas as pd
import re
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from docx import Document
from PIL import Image, ImageOps, ImageEnhance, ImageFilter
import pytesseract

# If Tesseract isn't on your system PATH (common on Windows), uncomment
# the line below and set it to wherever you installed Tesseract:
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\egwuo\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
# Regex patterns used to find emails and phone numbers within OCR'd text.
EMAIL_PATTERN = re.compile(r"[\w\.\-+]+@[\w\.\-]+\.\w+")
PHONE_PATTERN = re.compile(r"(\+?\d[\d\-\.\(\)\s]{7,}\d)")


def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    """
    Cleans up a photo to make OCR more accurate.

    Real-world phone photos often have uneven lighting, blur, or are
    too small/low-res for OCR to read reliably. These steps generally
    help a lot:
      1. Convert to grayscale (color isn't useful for reading text)
      2. Upscale if the image is small (OCR struggles with tiny text)
      3. Boost contrast (makes text edges stand out more clearly)
      4. Sharpen (helps with slight blur from handheld photos)
      5. Convert to pure black/white (binarization) - this is the
         format OCR engines are most reliably trained on
    """
    # Step 1: Grayscale
    image = image.convert("L")

    # Step 2: Upscale small images (helps OCR read small/distant text)
    # Tesseract generally performs best when text height is at least
    # ~20-30 pixels, so we scale up smaller images.
    min_width = 1500
    if image.width < min_width:
        scale_factor = min_width / image.width
        new_size = (int(image.width * scale_factor), int(image.height * scale_factor))
        image = image.resize(new_size, Image.LANCZOS)

    # Step 3: Boost contrast
    contrast_enhancer = ImageEnhance.Contrast(image)
    image = contrast_enhancer.enhance(2.0)  # 1.0 = no change, higher = more contrast

    # Step 4: Sharpen (helps counter slight motion blur/out-of-focus shots)
    image = image.filter(ImageFilter.SHARPEN)

    # Step 5: Binarize - convert to pure black & white using a threshold.
    # Any pixel darker than the threshold becomes black, lighter becomes white.
    threshold = 150
    image = image.point(lambda p: 255 if p > threshold else 0)

    return image


def extract_table_from_image(file_path: str) -> tuple[pd.DataFrame, list[str]]:
    """
    Runs OCR on a photo to extract Name/Phone/Email data.

    Since photos don't have neat columns like a spreadsheet, this works
    line-by-line:
      1. OCR reads all text from the image.
      2. For each line, we look for an email address and a phone number
         using pattern matching.
      3. Whatever text remains on that line (after removing the email
         and phone) is treated as the Name.

    Returns a tuple of:
      - a DataFrame with columns Name, Phone, Email (for lines that
        were successfully parsed)
      - a list of "unparsed" lines that didn't clearly contain a usable
        row, so the user can review/fix them manually.

    Works best with: printed text, good lighting, minimal skew/blur.
    Less reliable with: handwriting, blurry or angled photos.
    """
    image = Image.open(file_path)

    # Clean up the image first to improve OCR accuracy
    image = preprocess_image_for_ocr(image)

    # Run OCR — this returns one big block of text, with line breaks
    # roughly matching the lines in the photo.
    raw_text = pytesseract.image_to_string(image)

    rows = []
    unparsed_lines = []

    for line in raw_text.splitlines():
        line = line.strip()

        # Skip blank lines (common in OCR output)
        if not line:
            continue

        email_match = EMAIL_PATTERN.search(line)
        phone_match = PHONE_PATTERN.search(line)

        # We only count this as a usable row if we found AT LEAST
        # an email or a phone number — otherwise it's probably a
        # heading, stray text, or an OCR misread.
        if not email_match and not phone_match:
            unparsed_lines.append(line)
            continue

        email = email_match.group(0) if email_match else ""
        phone = phone_match.group(0) if phone_match else ""

        # Remove the matched email/phone from the line, whatever's
        # left over is treated as the Name.
        name = line
        if email_match:
            name = name.replace(email_match.group(0), "")
        if phone_match:
            name = name.replace(phone_match.group(0), "")
        name = name.strip(" ,-|\t")

        rows.append({"Name": name, "Phone": phone, "Email": email})

    if len(rows) == 0:
        raise ValueError(
            "No usable contact rows could be found in this photo. "
            "This usually means the text wasn't clear enough to read "
            "(try a sharper, well-lit, non-blurry photo) or the photo "
            "doesn't contain a list with emails or phone numbers."
        )

    df = pd.DataFrame(rows)
    return df, unparsed_lines


def extract_table_from_docx(file_path: str) -> pd.DataFrame:
    """
    Opens a .docx file and extracts the first table found inside it
    into a pandas DataFrame.

    Assumes the first row of the table is the header row
    (e.g. Name, Phone, Email).

    Raises a ValueError with a clear message if no table is found,
    so the web app can show a helpful error instead of crashing.
    """
    doc = Document(file_path)

    if len(doc.tables) == 0:
        raise ValueError(
            "No table was found in this Word document. "
            "NeatSheet needs the data to be in a table format "
            "(with columns like Name, Phone, Email)."
        )

    # Use the first table in the document.
    # (If there are multiple tables, we only use the first one for now.)
    table = doc.tables[0]

    # Row 0 is treated as the header (column names)
    header = [cell.text.strip() for cell in table.rows[0].cells]

    # Every row after that is actual data
    data_rows = []
    for row in table.rows[1:]:
        data_rows.append([cell.text.strip() for cell in row.cells])

    df = pd.DataFrame(data_rows, columns=header)
    return df


def clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """
    Takes a raw pandas DataFrame and returns a cleaned version:
    - Title Case names
    - Digits-only phone numbers
    - Duplicate emails removed
    """

    # --- Clean Name column ---
    if "Name" in df.columns:
        df["Name"] = df["Name"].astype(str).str.strip().str.title()

    # --- Clean Phone column ---
    if "Phone" in df.columns:
        df["Phone"] = (
            df["Phone"]
            .astype(str)
            .str.replace(r"[^0-9]", "", regex=True)
        )

    # --- Remove duplicate rows based on Email ---
    email_col = None
    for col in df.columns:
        if col.strip().lower() in ("email", "e-mail", "email address"):
            email_col = col
            break

    if email_col:
        df[email_col] = df[email_col].astype(str).str.strip().str.lower()
        df = df.drop_duplicates(subset=email_col, keep="first")

    return df


def save_clean_excel(df: pd.DataFrame, output_path: str, unparsed_lines: list[str] | None = None) -> None:
    """
    Saves a cleaned DataFrame to a nicely formatted Excel file
    at the given output_path.

    If unparsed_lines is provided (used for photo/OCR uploads), it adds
    a second sheet called "Needs Review" listing any lines from the
    photo that couldn't be confidently turned into a row, so the user
    can check/fix them manually.
    """
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Clean Data")
        worksheet = writer.sheets["Clean Data"]

        # Style the header row
        header_fill = PatternFill(start_color="305496", end_color="305496", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        header_alignment = Alignment(horizontal="center", vertical="center")

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = header_alignment

        # Auto-fit column widths
        for i, column in enumerate(df.columns, start=1):
            max_length = max(
                df[column].astype(str).map(len).max(),
                len(str(column))
            )
            col_letter = get_column_letter(i)
            worksheet.column_dimensions[col_letter].width = max_length + 4

        # Freeze header row
        worksheet.freeze_panes = "A2"

        # If there are unparsed lines (only happens with photo uploads),
        # add a second sheet so the user can see what needs manual review.
        if unparsed_lines:
            review_df = pd.DataFrame({"Unparsed text from photo": unparsed_lines})
            review_df.to_excel(writer, index=False, sheet_name="Needs Review")

            review_sheet = writer.sheets["Needs Review"]
            for cell in review_sheet[1]:
                cell.fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
                cell.font = Font(bold=True, color="FFFFFF")

            review_sheet.column_dimensions["A"].width = 60
